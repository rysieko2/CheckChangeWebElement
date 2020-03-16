# import time
from docx import Document

pathFolder = "/home/krzys/"
pathFile = "/home/krzys/Pattern-of-application.docx"
newPathFile = "/home/krzys/New-Application.docx"


def set_docx_edit_path_win():
    global pathFolder
    global pathFile
    global newPathFile
    pathFolder = "F:/"
    pathFile = "F:/Pattern-of-application.docx"
    newPathFile = "F:/New-Application.docx"


def changeWordInDocx(oldWord1, newWord1, oldWord2, newWord2, oldWord3, newWord3):
    document = Document(pathFile)
    for paragraph in document.paragraphs:
        if oldWord1 in paragraph.text:
            paragraph.text = paragraph.text.replace(oldWord1, newWord1)
        if oldWord2 in paragraph.text:
            paragraph.text = paragraph.text.replace(oldWord2, newWord2)
        if oldWord3 in paragraph.text:
            paragraph.text = paragraph.text.replace(oldWord3, newWord3)
    document.save(newPathFile)


def changeWordInDocxLinux(oldWord, newWord):
    document = pathFile.getDocument()
    search = document.createSearchDescriptor()
    search.SearchString = oldWord
    search.SearchAll = True
    search.SearchWords = True
    search.SearchCaseSensitive = False
    selsFound = document.findAll(search)
    if selsFound.getCount() == 0:
        return
    for selIndex in range(0, selsFound.getCount()):
        selFound = selsFound.getByIndex(selIndex)
        selFound.setString(newWord)

